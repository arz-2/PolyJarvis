#!/usr/bin/env python3
"""
Build paper/response_to_reviewers.docx from a structured scaffold.

Format mirrors paper/AMGPT_Reviewer_Responses.docx:
  "Report of Reviewer #N" header -> per comment: quoted comment, Response:, Action:.

This is the BULLET-SCAFFOLD pass (per user): Response/Action are short bullet lists of
intended reply points and planned changes, NOT finished prose. Each entry carries a
Status (Addressed / Partially / Deferred). Sim-dependent replies use [NUMBERS PENDING ...].
Regenerate after scaffolds become prose and statuses flip.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- status colors -------------------------------------------------------
STATUS_COLOR = {
    "Addressed":  RGBColor(0x1F, 0x77, 0x1F),   # green
    "Partially":  RGBColor(0xB8, 0x6A, 0x00),   # amber
    "Deferred":   RGBColor(0xB0, 0x00, 0x00),   # red
}

# ---- content: (id, comment, status, response_bullets, action_bullets) ----
R1_MAJOR = [
    ("Major 1 — LLM novelty / encoded vs. inferred",
     "Define what is encoded in tools/prompts/rules vs. what the LLM genuinely infers at runtime; otherwise hard to judge autonomy.",
     "Partially",
     ["We reframe the contribution as robust zero-shot orchestration and protocol adaptation over a heterogeneous, failure-prone toolchain -- not per-decision scientific novelty.",
      "We concede the point directly: on validated polymer classes the routine decisions (force field, charges, chain length, equilibration gate) are deterministic table lookups, not runtime inference. The planner transcribes encoded defaults verbatim for high-confidence classes (decision_policy.json confidence_gate).",
      "We add a three-tier Decision Architecture table (Tier 1 encoded schemas/constants/tool descriptions, Tier 2 structured table parameters, Tier 3 runtime LLM integration) anchored to decision_policy.json so the encoded/inferred boundary is an auditable artifact, not a claim.",
      "We localize the genuine LLM value to what a scripted pipeline cannot do: error recovery on unanticipated failures (Major 11) and off-table / low-confidence chemistry planning. Each Tier-3 example is screened by a single test -- could the if-then have been written in advance? If yes, it is reclassified as encoded.",
      "This is one concern shared with Major 7 (baseline), Major 11 (error recovery), and Minor 3 (language); we answer them with one shared evidence package (scripted-baseline ablation + error-recovery benchmark).",
      "Note: the formalized decision layer is part of this revision; we describe the current architecture and rest the claim on the re-run systems, not the original four-polymer results."],
     ["New Methods subsection 'Decision Architecture' (Tier 1/2/3 table anchored to decision_policy.json + honest caveat). [main, R1 Major 1 scaffold]",
      "Language audit: 'intelligent decision-making' -> 'LLM-orchestrated protocol adaptation'. [R1 Minor 3]",
      "Decision-provenance audit (% deterministic vs. inferred) as an honest bound -- context, not headline.",
      "Scripted-baseline ablation + error-recovery benchmark as primary novelty evidence. [shared with Major 7 / Major 11; NUMBERS PENDING revision-run logs]"]),
    ("Major 2 — Validation set too small",
     "Four polymers (three textbook) cannot establish generality across polarity, side-chain, tacticity, stiffness, H-bonding, etc.",
     "Partially",
     ["Expanding from 4 to 9 systems spanning 7 chemistries and 2 force fields.",
      "Added PLA, PVC, PSU, cis-PBD, PEEK (ester, vinyl halide, sulfone, diene, aromatic ketone).",
      "Failures will be reported honestly, not swapped for easier systems.",
      "Adding an explicit scope-exclusion statement (no semicrystalline / charged / copolymer)."],
     ["Expand benchmark + add scope-exclusion paragraph in Limitations. [main, R1 Major 2 scaffold]",
      "[NUMBERS PENDING 9-system x 5-replicate runs]"]),
    ("Major 3 — Pass/fail accounting not convincing",
     "The 5/8 framing may overstate; present strict and inclusive scorecards and avoid implying broad accuracy.",
     "Partially",
     ["Will present two explicit scorecards: strict (experiment-only) and inclusive (MD-lit accepted).",
      "PE Tg stated as excluded (ambiguous reference), not counted as a pass.",
      "PE density failure and the failing strict-Tg cases stated explicitly."],
     ["Two-row scorecard in Validation Summary. [main, R1 Major 3 scaffold]",
      "Abstract/conclusion lead with strict numbers. [NUMBERS PENDING]"]),
    ("Major 4 — PE density failure serious",
     "+25% density error indicates a major protocol failure; provide a post-mortem, rerun PE, and show autonomous detection/repair.",
     "Partially",
     ["Agree this is central, not peripheral.",
      "Adding a root-cause post-mortem (over-densification / kinetic trapping evidence: MSD, R_ee).",
      "Re-running PE with extended-equilibration corrected protocol; will show corrected vs. original.",
      "Proposing an over-densification auto-detect heuristic (final rho / crystalline rho > 0.95)."],
     ["PE post-mortem + auto-detect discussion. [main, R1 Major 4 / R2 C2 scaffold]",
      "[NUMBERS PENDING corrected PE re-run]"]),
    ("Major 5 — Tg cooling-rate / fitting artifacts",
     "Show Tg estimates are stable to T-range, step size, equilibration length, fit breakpoints; report replicate means and uncertainties.",
     "Partially",
     ["Will report replicate mean +/- std as primary (not best replicate).",
      "Adding a sensitivity study: fit method, cooling schedule, T-grid, equilibration duration.",
      "Full run-to-run spread shown rather than the favorable trajectory."],
     ["Tg table -> mean+/-std; SI sensitivity sub-table. [main R1 Major 5/6, SI scaffolds]",
      "[NUMBERS PENDING 5-replicate runs]"]),
    ("Major 6 — 'Best replicate' biases validation",
     "Judge by expected performance over independent runs; make primary validation ensemble-based.",
     "Partially",
     ["Primary validation moved to ensemble mean +/- std across 5 replicates.",
      "Best-replicate values demoted to secondary/SI discussion."],
     ["Replicate statistics as primary in all property tables/figures.",
      "[NUMBERS PENDING 5-replicate runs]"]),
    ("Major 7 — Baseline comparison needed",
     "Compare against stock RadonPy defaults / expert-curated / non-agent scripted workflows.",
     "Partially",
     ["Disclose the RadonPy->EMC builder migration explicitly (16/21 classes now EMC); label the builder per system.",
      "Primary baseline isolates the agent, not the builder: same EMC/LAMMPS backend, scripted-default vs agent, on 2-3 systems incl. a failure case.",
      "A RadonPy-manual vs EMC-agent comparison would confound the builder swap with the agent; we avoid it.",
      "Keep stock-RadonPy-defaults on the original 4 as a continuity/effort reference, not the isolation control.",
      "Frame value as reduced user effort + error recovery, not necessarily higher raw accuracy; acknowledge where defaults already match."],
     ["EMC-disclosure in Methods + same-backend baseline subsection. [main, R1 Major 7 scaffold]",
      "[NUMBERS PENDING scripted-vs-agent baseline runs]"]),
    ("Major 8 — FF / charge selection justification",
     "Provide evidence choices are appropriate; include FF/charge sensitivity for 1-2 systems where predictions fail.",
     "Partially",
     ["Add a per-class FF-selection rationale table (force field, charge method, literature citations) from the encoded polymer rules.",
      "Clarify the GAFF2/Gasteiger point: Gasteiger is used only on nonpolar united-atom classes (|q|<0.1 e, negligible Coulomb; Afzal2021); polar systems use PCFF/OPLS-AA with RESP or AM1-BCC.",
      "Still run an FF/charge sensitivity test on a failed case (PE density or aPS Tg); a negative result is still informative."],
     ["FF-rationale table (writable now) + FF-sensitivity subsection. [main, R1 Major 8 scaffold]",
      "[NUMBERS PENDING FF sensitivity runs]"]),
    ("Major 9 — Bulk modulus statistically fragile",
     "Report autocorrelation times, effective sample sizes, block-averaged uncertainties, barostat sensitivity.",
     "Partially",
     ["Tooling now emits tau_eff, N_eff, and block-averaged SEM for K.",
      "Adding a barostat tau_P sensitivity study (PMMA, tau_P = 250/1000/4000 fs).",
      "Compare tau_P-insensitive K_Born (NVT) against K_fluc; clarify K_T vs K_S references."],
     ["SI S7 columns + barostat table; main-text robustness note. [R1 Major 9 scaffolds]",
      "[NUMBERS PENDING barostat-sensitivity runs]"]),
    ("Major 10 — Structural validation too qualitative",
     "Add Rg distributions, characteristic ratio, MSD, relaxation indicators, density homogeneity, esp. for trapped PE.",
     "Partially",
     ["Updated equil-checker now produces Rg distributions, MSD plateau, density homogeneity, and C(t) KWW relaxation time.",
      "For PE, will show explicit chain-relaxation vs. kinetic-trapping evidence."],
     ["Expanded structural diagnostics subsection. [main, R1 Major 10 scaffold]",
      "[NUMBERS PENDING revision runs through updated checker]"]),
    ("Major 11 — Error recovery evaluation",
     "Provide a reproducible error-recovery benchmark: attempts/failures, human guidance, pre-programmed vs. inferred, success rates, logs.",
     "Deferred",
     ["Will assemble a per-incident benchmark table from the completed revision-run logs (not back-filled from exploratory runs).",
      "Columns: error class, recovery action, pre-scripted vs. Tier-3, #retries incl. failures, outcome, human guidance.",
      "Report an overall recovery success rate."],
     ["Reframe SI error table as benchmark (structure only this pass). [SI, R1 Major 11 scaffold]",
      "[NUMBERS PENDING revision-run logs]"]),
    ("Major 12 — Reproducibility / data availability",
     "Workflow traces, prompts, tool calls, decision logs, input scripts, structures, seeds, notebooks must be openly archived.",
     "Partially",
     ["Agree provenance is core, not supplementary.",
      "Will deposit an archived release (DOI) with prompts, tool schemas, input scripts, seeds, final structures, notebooks, raw Tg-fit data.",
      "Availability statement updated to cite the archived-release DOI."],
     ["Rewrite Data & Software Availability. [main, R1 Major 12 scaffold]",
      "[PENDING Zenodo DOI + release tag]"]),
]

R1_MINOR = [
    ("Minor 1 — Title precision", "Title overclaims scope.", "Partially",
     ["Will make title scope-accurate (amorphous homopolymers; orchestrated agent)."],
     ["Title revision candidate provided. [main, R1 Minor 1 scaffold]"]),
    ("Minor 2 — Abstract balance", "Distinguish successful/failed/excluded cases; note only PMMA passes strict Tg and PE density fails.", "Partially",
     ["Restructure abstract to lead with strict scorecard, PE density failure, failing Tg cases."],
     ["Abstract scaffold. [main, R1 Minor 2] [NUMBERS PENDING]"]),
    ("Minor 3 — 'Intelligent decision-making'", "Use more defensible language.", "Partially",
     ["Replace with 'LLM-integrated protocol adaptation' / tier-appropriate phrasing."],
     ["Language audit. [main, R1 Major 1 / Minor 3 scaffold]"]),
    ("Minor 4 — MCP role", "Explain what MCP adds beyond ordinary API calls/wrappers.", "Addressed",
     ["Add typed schemas, multi-turn invocation, lifecycle, error propagation; runtime order vs. fixed DAG."],
     ["Main-text + SI S1 scaffolds. [R1 Minor 4]"]),
    ("Minor 5 — Conversation figure generic", "Replace placeholder with a real trace incl. a decision and recovery.", "Deferred",
     ["Will use the real PMMA4 D-05 trace (tool call + reasoning + Tier-3 decision), color-annotated."],
     ["Figure deferred; caption/intent scaffolded. [main, R1 Minor 5]"]),
    ("Minor 6 — Benchmark table mixes refs", "Visually separate experimental vs. MD-literature refs.", "Partially",
     ["Add shading/grouping for (A) experimental, (B) MD-lit, (C) excluded; PE/PEG basis obvious at a glance."],
     ["Benchmark table visual scaffold. [main, R1 Minor 6]"]),
    ("Minor 7 — PE density plot prominence", "Make PE failure visually prominent.", "Deferred",
     ["Render PE distinctly (red/hatch/arrow); call it out in caption."],
     ["Figure deferred; intent scaffolded. [main, R1 Minor 7] [NUMBERS PENDING PE re-run]"]),
    ("Minor 8 — Tg figure replicates", "Show all replicates / uncertainty in the main figure.", "Deferred",
     ["Main Tg figure: scatter of all replicates with mean+/-std bar; per-run curves in SI."],
     ["Figure deferred; intent scaffolded. [main, R1 Minor 8] [NUMBERS PENDING]"]),
    ("Minor 9 — Why 10 chains", "Justify or test the 10-chain choice.", "Partially",
     ["Cite RadonPy/GAFF2 convergence at ~10 chains; note finite-size caveats; offer 5/10/20 sensitivity."],
     ["SI S6 justification scaffold. [R1 Minor 9]"]),
    ("Minor 10 — Tacticity", "Describe how tacticity was generated, verified, varied for aPS/PMMA.", "Partially",
     ["Describe RadonPy stereochemistry generation and per-chain distribution; offer atactic-vs-syndiotactic PMMA sensitivity."],
     ["New SI Tacticity subsection. [R1 Minor 10] [NUMBERS PENDING if sensitivity run]"]),
    ("Minor 11 — Soften cooling-rate attribution", "FF, chain length, equilibration, fitting can also contribute.", "Partially",
     ["Reword to 'consistent with cooling-rate bias, though FF/chain-length/equilibration/fitting cannot be ruled out'."],
     ["Cooling-rate softening scaffold (abstract/results/limitations). [main, R1 Minor 11]"]),
    ("Minor 12 — Software availability", "Need exact version, scripts, prompts, schemas, seeds, archived release.", "Partially",
     ["Covered by the archived-release DOI (see Major 12)."],
     ["Availability rewrite. [main, R1 Major 12 / Minor 12] [PENDING DOI]"]),
    ("Minor 13 — Computational cost", "Report wall time, GPU hours, approx cost, retry overhead.", "Partially",
     ["Add main-text cost summary + SI S12 columns (GPU-h, API calls/cost, retry overhead)."],
     ["Main + SI cost scaffolds. [R1 Minor 13] [NUMBERS PENDING revision runs]"]),
    ("Minor 14 — MD agreement != validation", "Agreement with prior MD may share biases; phrase cautiously.", "Addressed",
     ["Add qualifier: methodological consistency, not independent experimental validation; shared systematic biases possible."],
     ["MD-agreement qualifier scaffold + audit list. [main, R1 Minor 14]"]),
    ("Minor 15 — Conclusion restraint", "Emphasize proof of concept, failure modes, need for broader validation.", "Partially",
     ["Rewrite conclusion to lead with proof-of-concept and identified failure modes; drop production-ready phrasing."],
     ["Conclusion scaffold. [main, R1 Minor 15] [NUMBERS PENDING]"]),
]

R2 = [
    ("Comment 1 — Comparison with prior frameworks",
     "Table 1 reviews prior frameworks but does not discuss how PolyJarvis compares; add discussion.",
     "Addressed",
     ["Add a paragraph after Table 1 comparing scope, autonomy, pipeline integration, property coverage, validation scale.",
      "Be honest where prior tools match or exceed PolyJarvis."],
     ["Table 1 comparison-paragraph scaffold. [main, R2 C1]"]),
    ("Comment 2 — PE density / autonomous detection",
     "Density overestimated >25%; have authors tried to improve the protocol and detect equilibration failure autonomously?",
     "Partially",
     ["Re-running PE with corrected protocol; report corrected vs. original.",
      "Discuss an autonomous over-densification detection heuristic."],
     ["PE post-mortem + auto-detect scaffold. [main, R1 Major 4 / R2 C2] [NUMBERS PENDING PE re-run]"]),
    ("Comment 3 — Origin of errors",
     "Clarify whether discrepancies arise from workflow, equilibration, force field, or sampling.",
     "Partially",
     ["Add a 'Sources of Prediction Error' subsection categorizing each case (agent/protocol, FF, cooling-rate, sampling).",
      "Anchor to the FF-sensitivity results."],
     ["New Sources-of-Prediction-Error subsection. [main, R2 C3] [NUMBERS PENDING FF test]"]),
]


def add_heading_rule(doc, text):
    p = doc.add_paragraph()
    p.add_run("-" * 70)
    h = doc.add_paragraph()
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.add_run("-" * 70)


def add_entry(doc, cid, comment, status, responses, actions):
    # comment
    pc = doc.add_paragraph()
    rc = pc.add_run(f"{cid}")
    rc.bold = True
    pq = doc.add_paragraph()
    pq.add_run(comment).italic = True
    # status
    ps = doc.add_paragraph()
    ps.add_run("Status: ").bold = True
    rs = ps.add_run(status)
    rs.bold = True
    rs.font.color.rgb = STATUS_COLOR.get(status, RGBColor(0, 0, 0))
    # response
    doc.add_paragraph().add_run("Response:").bold = True
    for b in responses:
        doc.add_paragraph(b, style="List Bullet")
    # action
    doc.add_paragraph().add_run("Action:").bold = True
    for b in actions:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph("")  # spacer


def main():
    doc = Document()
    title = doc.add_paragraph()
    rt = title.add_run("Response to the reviewers' comments")
    rt.bold = True
    rt.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph()
    rn = note.add_run(
        "Manuscript ct-2026-00736q. NOTE: this is a SCAFFOLD draft — Response/Action are "
        "intended reply points (bullets), not final prose; items tagged [NUMBERS PENDING] await "
        "in-progress simulations. Regenerate via build_response_docx.py as items are finalized.")
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_heading_rule(doc, "Report of Reviewer #1 — ct-2026-00736q")
    doc.add_paragraph().add_run("Major Comments").bold = True
    for e in R1_MAJOR:
        add_entry(doc, *e)
    doc.add_paragraph().add_run("Minor Comments").bold = True
    for e in R1_MINOR:
        add_entry(doc, *e)

    add_heading_rule(doc, "Report of Reviewer #2 — ct-2026-00736q")
    for e in R2:
        add_entry(doc, *e)

    out = "response_to_reviewers.docx"
    doc.save(out)
    n = len(R1_MAJOR) + len(R1_MINOR) + len(R2)
    print(f"Wrote {out} with {n} comment entries "
          f"(R1 major {len(R1_MAJOR)}, R1 minor {len(R1_MINOR)}, R2 {len(R2)}).")


if __name__ == "__main__":
    main()
